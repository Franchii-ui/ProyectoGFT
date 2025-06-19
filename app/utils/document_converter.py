import os
import json
from pathlib import Path
from docx import Document
from fpdf import FPDF
from jinja2 import Template
import logging

from app.utils import document_converter

logger = logging.getLogger(__name__)

def text_to_docx(text: str, output_path: Path) -> Path:
    """Convert plain text to DOCX format"""
    try:
        doc = Document()
        doc.add_heading('Transcription', level=1)
        paragraphs = text.split('\n\n')
        for para in paragraphs:
            if para.strip():
                doc.add_paragraph(para.strip())
        doc.save(output_path)
        return output_path
    except Exception as e:
        logger.error(f"Error creating DOCX document: {str(e)}")
        raise

def text_to_pdf(text: str, output_path: Path) -> Path:
    """Convert plain text to PDF format"""
    try:
        print(f"DEBUG: text_to_pdf received text: {repr(text[:100])}")  # Show first 100 chars
        pdf = FPDF()
        pdf.add_page()
        pdf.set_font("Arial", 'B', 16)
        pdf.cell(200, 10, txt="Transcription", ln=True, align='C')
        pdf.ln(10)
        pdf.set_font("Arial", size=12)
        paragraphs = text.split('\n\n')
        for para in paragraphs:
            if para.strip():
                pdf.multi_cell(0, 10, txt=para.strip())
                pdf.ln(5)
        pdf.output(output_path)
        return output_path
    except Exception as e:
        logger.error(f"Error creating PDF document: {str(e)}")
        raise

def text_to_html(text: str, output_path: Path) -> Path:
    """Convert plain text to HTML format"""
    try:
        html_template = """
        <!DOCTYPE html>
        <html>
        <head>
            <title>Transcription</title>
            <style>
                body { font-family: Arial, sans-serif; line-height: 1.6; max-width: 800px; margin: 0 auto; padding: 20px; }
                h1 { color: #333; text-align: center; }
                p { margin-bottom: 16px; }
            </style>
        </head>
        <body>
            <h1>Transcription</h1>
            {% for paragraph in paragraphs %}
                <p>{{ paragraph }}</p>
            {% endfor %}
        </body>
        </html>
        """
        template = Template(html_template)
        paragraphs = [p.strip() for p in text.split('\n\n') if p.strip()]
        html_content = template.render(paragraphs=paragraphs)
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(html_content)
        return output_path
    except Exception as e:
        logger.error(f"Error creating HTML document: {str(e)}")
        raise