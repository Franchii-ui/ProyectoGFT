from fastapi import APIRouter, HTTPException, Request, Form, Depends, Response
from fastapi.responses import FileResponse, JSONResponse
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy.future import select
from pathlib import Path
from typing import Optional
from app.utils import document_converter
import app.config as config
import logging
from enum import Enum
from app.db import get_db
from app.models.transcription import Transcription
from app.models.user import User
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from io import BytesIO
import re
import requests
import os
from openai import OpenAI
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

logger = logging.getLogger(__name__)
router = APIRouter()

class DocumentFormat(str, Enum):
    TXT = "txt"
    DOCX = "docx"
    PDF = "pdf"
    HTML = "html"



def chunk_text(text, max_length=2000):
    """Split text into chunks of max_length (in characters)."""
    paragraphs = text.split('\n')
    chunks = []
    current_chunk = ""
    for para in paragraphs:
        if len(current_chunk) + len(para) + 1 > max_length:
            chunks.append(current_chunk)
            current_chunk = para
        else:
            if current_chunk:
                current_chunk += "\n" + para
            else:
                current_chunk = para
    if current_chunk:
        chunks.append(current_chunk)
    return chunks

def humanize_with_gpt35(text, language="es"):
    api_key = os.getenv("OPENAI_API_KEY")
    if not api_key:
        raise RuntimeError("OPENAI_API_KEY not set in environment variables")
    client = OpenAI(api_key=api_key)
    prompt = (
        f"Corrige solamente la gramática y puntuación en el siguiente texto en {language}. "
        "NO añadas texto propio. NO resumas. NO cambies palabras. NO agregues comentarios. "
        "SOLO corrige errores gramaticales, puntuación y capitalización:\n\n"
        + text
    )
    response = client.chat.completions.create(
        model="gpt-3.5-turbo",
        messages=[
            {"role": "system", "content": "Eres un corrector ortográfico profesional."},
            {"role": "user", "content": prompt}
        ],
        max_tokens=2048,
        temperature=0.0,
    )
    return response.choices[0].message.content.strip()

def section_and_title_with_gpt35(text, language="es"):
    api_key = os.getenv("OPENAI_API_KEY")
    if not api_key:
        raise RuntimeError("OPENAI_API_KEY not set in environment variables")
    client = OpenAI(api_key=api_key)
    prompt = (
        f"Divide el siguiente texto en secciones lógicas. "
        "Para cada sección, genera un título breve y descriptivo (máximo 8 palabras), y coloca el texto de esa sección debajo del título. "
        "Devuelve el resultado en el siguiente formato:\n\n"
        "[TÍTULO DE LA SECCIÓN 1]\nTexto de la sección 1...\n\n"
        "[TÍTULO DE LA SECCIÓN 2]\nTexto de la sección 2...\n\n"
        "No agregues comentarios ni texto adicional.\n\n"
        + text
    )
    response = client.chat.completions.create(
        model="gpt-3.5-turbo",
        messages=[
            {"role": "system", "content": "Eres un editor profesional de textos."},
            {"role": "user", "content": prompt}
        ],
        max_tokens=2048,
        temperature=0.3,
    )
    return response.choices[0].message.content.strip()

def add_toc(paragraph):
    fldSimple = OxmlElement('w:fldSimple')
    fldSimple.set(qn('w:instr'), 'TOC \\o "1-3" \\h \\z \\u')
    paragraph._p.append(fldSimple)

@router.get("/download/{file_id}")
async def download_transcription(file_id: str):
    file_path = config.RESULTS_DIR / f"{file_id}.txt"
    if not file_path.exists():
        raise HTTPException(status_code=404, detail="Transcription not found")
    return FileResponse(
        path=str(file_path),
        filename=f"transcription_{file_id}.txt",
        media_type="text/plain"
    )

@router.get("/export/{file_id}")
async def export_transcription(
    file_id: str,
    format: str,
    language: str = "es",  # <-- Default to Spanish
    db: AsyncSession = Depends(get_db)
):
    txt_path = config.RESULTS_DIR / f"{file_id}.txt"
    if not txt_path.exists():
        raise HTTPException(status_code=404, detail="Transcription not found")
    with open(txt_path, 'r') as f:
        transcription_text = f.read()

    # Clean/humanize the transcription here
    corrected_text = humanize_with_gpt35(transcription_text, language)

    # If empty or too short, use original
    if not corrected_text or len(corrected_text) < len(transcription_text) * 0.8:
        print("Warning: Using original text - AI result seems incomplete")
        corrected_text = transcription_text

    print("Transcription after Llama3:", corrected_text)

    if not corrected_text.strip():
        corrected_text = "No se pudo procesar la transcripción."

    output_filename = f"{file_id}.{format}"
    output_path = config.RESULTS_DIR / output_filename

    if format == DocumentFormat.DOCX:
        doc = Document()

        # --- COVER PAGE (Section 1, margins 0) ---
        section = doc.sections[0]
        section.top_margin = Inches(0)
        section.bottom_margin = Inches(0)
        section.left_margin = Inches(0)
        section.right_margin = Inches(0)

        # Add cover image (almost full page, leave a bit of space for section break)
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture('images/cover.png', width=Inches(8.27), height=Inches(10.5))  # Slightly less than full A4

        # --- MAIN CONTENT (Section 2, standard margins) ---
        main_section = doc.sections[1]
        main_section.top_margin = Inches(1)
        main_section.bottom_margin = Inches(1)
        main_section.left_margin = Inches(1)
        main_section.right_margin = Inches(1)

        # Add header logo ONLY to section 2
        header = main_section.header
        header_paragraph = header.paragraphs[0]
        header_run = header_paragraph.add_run()
        header_run.add_picture('images/GFT_Logo_RGB.png', width=Inches(1.0))

        # Table of Contents
        doc.add_paragraph('Tabla de Contenidos', style='Heading 1')
        toc_paragraph = doc.add_paragraph()
        add_toc(toc_paragraph)
        doc.add_page_break()

        # --- Sectioned Content ---
        sectioned_text = section_and_title_with_gpt35(corrected_text, language)

        # Parse sections: [TITLE]\nText...
        pattern = re.compile(r"\[(.*?)\]\n(.*?)(?=\n\[|$)", re.DOTALL)
        matches = pattern.findall(sectioned_text)

        if not matches:
            # fallback: just add the text as a single section
            doc.add_heading('Transcripción de Video', level=1)
            paragraph = doc.add_paragraph(sectioned_text or "Sin texto")
            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            for run in paragraph.runs:
                run.font.name = 'Noto Sans'
                run.font.size = Pt(12)
        else:
            for title, content in matches:
                heading = doc.add_heading(title.strip(), level=1)
                heading.runs[0].font.name = 'Arial'
                heading.runs[0].font.size = Pt(18)
                paragraph = doc.add_paragraph(content.strip())
                paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                for run in paragraph.runs:
                    run.font.name = 'Noto Sans'
                    run.font.size = Pt(12)

        file_stream = BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)
        headers = {
            "Content-Disposition": f"attachment; filename=gft_transcription.docx"
        }
        return Response(
            file_stream.read(),
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers=headers
        )
    elif format == DocumentFormat.PDF:
        output_path = document_converter.text_to_pdf(corrected_text, output_path)
        media_type = "application/pdf"
        download_name = "gft_transcription.pdf"
    elif format == DocumentFormat.HTML:
        output_path = document_converter.text_to_html(corrected_text, output_path)
        media_type = "text/html"
        download_name = "gft_transcription.html"
    else:
        output_path = txt_path
        media_type = "text/plain"
        download_name = "gft_transcription.txt"
    if format != DocumentFormat.DOCX:
        return FileResponse(
            path=str(output_path),
            filename=download_name,
            media_type=media_type
        )

@router.post("/edit/{file_id}")
async def edit_transcription(file_id: str, text: str = Form(...)):
    txt_path = config.RESULTS_DIR / f"{file_id}.txt"
    if not txt_path.exists():
        raise HTTPException(status_code=404, detail="Transcription not found")
    with open(txt_path, "w") as f:
        f.write(text)
    return JSONResponse(content={"success": True, "message": "Transcription updated."})

@router.post("/save/{file_id}")
async def save_transcription(file_id: str, request: Request):
    data = await request.json()
    text = data.get("text")
    if not text:
        raise HTTPException(status_code=400, detail="No text provided")
    txt_path = config.RESULTS_DIR / f"{file_id}.txt"
    with open(txt_path, "w") as f:
        f.write(text)
    return JSONResponse(content={"success": True, "message": "Transcription saved."})